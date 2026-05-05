from __future__ import annotations

import json
import os
import re
import shutil
import urllib.request
from datetime import datetime
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
except ImportError as exc:
    print("ERROR: openpyxl is not installed. Install it in project venv: pip install openpyxl")
    raise SystemExit(1) from exc


KB_ITEMS_HEADERS = [
    "id",
    "enabled",
    "category",
    "type",
    "title",
    "priority",
    "topic_keywords",
    "entity_keywords",
    "customer_need_keywords",
    "response_modes",
    "applies_to",
    "content",
    "operator_instruction",
    "template_hint",
    "source",
    "source_type",
    "source_ref",
    "source_quote",
    "source_date",
    "notes",
]

KB_DICTIONARY_HEADERS = ["entity_id", "type", "canonical_value", "aliases", "notes"]
KB_TYPES_HEADERS = ["type", "enabled", "description"]
KB_SETTINGS_HEADERS = ["key", "value", "description"]

DEFAULT_TYPES = [
    ["rule", "yes", "Правило обработки класса обращений."],
    ["faq", "yes", "Стабильный фрагмент ответа на частый вопрос."],
    ["operator_instruction", "yes", "Внутренняя инструкция оператору, не прямой текст клиенту."],
    ["response_template", "yes", "Подсказка для будущего Draft Builder, не готовый ответ."],
    ["regulation", "yes", "Внутреннее или внешнее правило, ограничивающее обработку."],
    ["historical_case", "no", "Зарезервировано на будущее. В MVP не использовать без явного решения."],
]

DEFAULT_SETTINGS = [
    ["min_score", "2", "Минимальный score для будущего Knowledge Retrieval MVP."],
    ["max_results", "3", "Максимальное количество найденных KB_ITEMS."],
    ["enabled_types", "rule;faq;operator_instruction;response_template;regulation", "Активные типы знаний."],
    ["no_result_behavior", "return_empty", "Если ничего не найдено, вернуть empty result и не выдумывать контекст."],
    ["source_format", "xlsx", "Knowledge Base хранится в одном XLSX workbook."],
]

DEFAULT_DICTIONARY = [
    ["brand_mooon", "brand", "mooon", "mooon; moon; мун; муун; mooon cinema; mooon.by", "Бренд mooon."],
    ["brand_silverscreen", "brand", "SilverScreen", "SilverScreen; Silver Screen; сильверскрин; сильвер скрин", "Бренд SilverScreen."],
    ["service_go2", "service", "go2.by", "go2; go2.by; го2; гоу2", "Виджет/сервис продажи билетов."],
    ["venue_dana_mall", "venue", "mooon в Dana Mall", "Dana Mall; Дана Молл; Дана; mooon Dana", "Кинотеатр mooon в Dana Mall."],
    ["venue_palazzo", "venue", "mooon в Palazzo", "Palazzo; Палаццо; mooon Palazzo", "Кинотеатр mooon в Palazzo."],
    ["venue_arenacity", "venue", "mooon в ArenaCity", "ArenaCity; АренаСити; Арена; mooon ArenaCity", "Кинотеатр mooon в ArenaCity."],
    ["venue_trinity", "venue", "mooon в Trinity Гродно", "Trinity; Тринити; Гродно; mooon Trinity", "Кинотеатр mooon в Trinity Гродно."],
]

ALLOWED_ITEM_TYPES = {"rule", "faq", "operator_instruction", "response_template", "regulation", "historical_case"}
ALLOWED_RESPONSE_MODES = {"answer", "ask_clarifying_question", "handoff_to_operator", "no_reply/ignore"}

STOPWORDS_RU = {
    "это", "как", "что", "для", "или", "при", "над", "под", "без", "если", "так", "его", "еще", "уже",
    "все", "она", "они", "оно", "где", "кто", "чем", "чего", "тогда", "также", "после", "перед",
    "можно", "нужно", "должен", "должна", "должно", "который", "которая", "которые", "этот", "эта", "эти",
    "клиент", "клиента", "оператор", "оператору", "обращение", "письмо", "ответ", "вопрос",
}


class HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style", "noscript", "svg"}:
            self.skip_depth += 1
        if tag in {"p", "br", "div", "section", "article", "li", "tr", "h1", "h2", "h3", "h4", "blockquote"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript", "svg"} and self.skip_depth > 0:
            self.skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if self.skip_depth:
            return
        text = data.strip()
        if text:
            self.parts.append(text)

    def get_text(self) -> str:
        text = " ".join(self.parts)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n\s+", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return clean_text(text)


def find_project_root() -> Path:
    current = Path(__file__).resolve()
    for parent in [current.parent, *current.parents]:
        if (parent / "SKILLS.md").exists() or (parent / "config").exists():
            return parent
    return Path.cwd()


PROJECT_ROOT = find_project_root()
KB_PATH = PROJECT_ROOT / "config" / "knowledge_base" / "knowledge_base.xlsx"
BACKUP_PATH = PROJECT_ROOT / "config" / "knowledge_base" / "knowledge_base.backup.xlsx"


def clean_text(text: str) -> str:
    text = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_text_file(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1251"):
        try:
            return clean_text(path.read_text(encoding=enc))
        except UnicodeDecodeError:
            continue
    raise RuntimeError(f"Cannot read text file: {path}")


def read_html_file(path: Path) -> str:
    html = read_text_file(path)
    parser = HTMLTextExtractor()
    parser.feed(html)
    return parser.get_text()


def read_pdf_file(path: Path) -> str:
    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise RuntimeError("PDF reading requires PyMuPDF. Install it in project venv: pip install pymupdf") from exc

    parts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return clean_text("\n\n".join(parts))


def read_docx_file(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("DOCX reading requires python-docx. Install it in project venv: pip install python-docx") from exc

    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return clean_text("\n".join(parts))


def read_xlsx_source_file(path: Path) -> str:
    wb = load_workbook(path, data_only=True, read_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            values = [str(v).strip() for v in row if v not in (None, "")]
            if values:
                parts.append(" | ".join(values))
    return clean_text("\n".join(parts))


def read_image_file(path: Path) -> str:
    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Image OCR requires pillow and pytesseract. Install them and make sure Tesseract OCR is available.") from exc

    text = pytesseract.image_to_string(Image.open(path), lang=os.getenv("TESSERACT_LANG", "rus+eng"))
    return clean_text(text)


def read_file(path: Path) -> str:
    if not path.exists():
        raise RuntimeError(f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".csv", ".log"}:
        return read_text_file(path)
    if suffix in {".html", ".htm"}:
        return read_html_file(path)
    if suffix == ".pdf":
        return read_pdf_file(path)
    if suffix == ".docx":
        return read_docx_file(path)
    if suffix in {".xlsx", ".xlsm"}:
        return read_xlsx_source_file(path)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
        return read_image_file(path)
    if suffix in {".mp4", ".mov", ".avi", ".mkv", ".webm"}:
        raise RuntimeError("Video files are not parsed by this runner. Provide a transcript/text summary as input.")

    return read_text_file(path)


def read_url(url: str) -> str:
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 KB update"})
    with urllib.request.urlopen(req, timeout=30) as resp:
        raw = resp.read().decode("utf-8", errors="replace")
    parser = HTMLTextExtractor()
    parser.feed(raw)
    return parser.get_text()


def infer_source_type(source_ref: str) -> str:
    if source_ref == "manual_text":
        return "manual_text"
    if re.match(r"^https?://", source_ref, flags=re.I):
        return "url"
    suffix = Path(source_ref).suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix in {".xlsx", ".xlsm"}:
        return "xlsx"
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
        return "image"
    if suffix in {".html", ".htm"}:
        return "html"
    if suffix:
        return suffix.lstrip(".")
    return "unknown"


def normalize(value: str) -> str:
    value = str(value or "").lower().replace("ё", "е")
    value = re.sub(r"[^a-zа-я0-9]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def split_keywords(value: str) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for part in re.split(r"[;\n,]+", str(value or "")):
        item = part.strip()
        key = item.lower()
        if item and key not in seen:
            seen.add(key)
            result.append(item)
    return result


def join_unique(*values: str) -> str:
    result: list[str] = []
    seen: set[str] = set()
    for value in values:
        for item in split_keywords(value):
            key = item.lower()
            if key not in seen:
                seen.add(key)
                result.append(item)
    return "; ".join(result)


def token_set(value: str) -> set[str]:
    return {x for x in normalize(value).split() if len(x) >= 4 and x not in STOPWORDS_RU}


def similarity(a: str, b: str) -> float:
    aa, bb = token_set(a), token_set(b)
    if not aa or not bb:
        return 0.0
    return len(aa & bb) / len(aa | bb)


def slugify(value: str, fallback: str = "kb_item") -> str:
    value = normalize(value)
    replacements = {
        "возврат": "refund",
        "вернуть": "refund",
        "деньги": "money",
        "билет": "ticket",
        "письмо": "email",
        "расписание": "schedule",
        "репертуар": "repertoire",
        "сеанс": "session",
        "оплата": "payment",
        "ошибка": "error",
        "доставка": "delivery",
        "почта": "email",
        "правило": "rule",
        "инструкция": "instruction",
        "регламент": "regulation",
    }
    for src, dst in replacements.items():
        value = value.replace(src, dst)
    value = re.sub(r"[^a-z0-9]+", "_", value)
    value = re.sub(r"_+", "_", value).strip("_")
    return value[:64] or fallback


def sheet_headers(ws: Any, required_headers: list[str]) -> list[str]:
    headers: list[str] = []
    for col_idx in range(1, max(ws.max_column, 1) + 1):
        value = ws.cell(row=1, column=col_idx).value
        if value is not None and str(value).strip():
            headers.append(str(value).strip())
    if not headers:
        headers = list(required_headers)
        for col_idx, header in enumerate(headers, start=1):
            ws.cell(row=1, column=col_idx).value = header
    return headers


def ensure_sheet(wb: Any, sheet_name: str, required_headers: list[str]) -> None:
    if sheet_name not in wb.sheetnames:
        ws = wb.create_sheet(sheet_name)
        ws.append(required_headers)
        return

    ws = wb[sheet_name]
    headers = sheet_headers(ws, required_headers)
    existing = set(headers)
    col_idx = len(headers) + 1
    for header in required_headers:
        if header not in existing:
            ws.cell(row=1, column=col_idx).value = header
            headers.append(header)
            existing.add(header)
            col_idx += 1


def seed_sheet(ws: Any, required_headers: list[str], rows: list[list[str]]) -> None:
    headers = sheet_headers(ws, required_headers)
    existing_keys = set()
    key_header = headers[0]
    key_col = headers.index(key_header) + 1
    for row_idx in range(2, ws.max_row + 1):
        value = ws.cell(row=row_idx, column=key_col).value
        if value not in (None, ""):
            existing_keys.add(str(value).strip().lower())

    for row in rows:
        key = str(row[0]).strip().lower()
        if key in existing_keys:
            continue
        data = {required_headers[i]: row[i] for i in range(min(len(required_headers), len(row)))}
        append_row(ws, headers, data)
        existing_keys.add(key)


def ensure_workbook() -> Any:
    KB_PATH.parent.mkdir(parents=True, exist_ok=True)
    if KB_PATH.exists():
        wb = load_workbook(KB_PATH)
    else:
        wb = Workbook()
        wb.remove(wb.active)

    ensure_sheet(wb, "KB_ITEMS", KB_ITEMS_HEADERS)
    ensure_sheet(wb, "KB_DICTIONARY", KB_DICTIONARY_HEADERS)
    ensure_sheet(wb, "KB_TYPES", KB_TYPES_HEADERS)
    ensure_sheet(wb, "KB_SETTINGS", KB_SETTINGS_HEADERS)

    seed_sheet(wb["KB_TYPES"], KB_TYPES_HEADERS, DEFAULT_TYPES)
    seed_sheet(wb["KB_SETTINGS"], KB_SETTINGS_HEADERS, DEFAULT_SETTINGS)
    seed_sheet(wb["KB_DICTIONARY"], KB_DICTIONARY_HEADERS, DEFAULT_DICTIONARY)
    style_workbook(wb)
    return wb


def style_workbook(wb: Any) -> None:
    fill = PatternFill("solid", fgColor="1F2937")
    font = Font(color="FFFFFF", bold=True)
    default_widths = {
        "id": 28,
        "enabled": 10,
        "category": 22,
        "type": 22,
        "title": 38,
        "priority": 10,
        "topic_keywords": 36,
        "entity_keywords": 34,
        "customer_need_keywords": 42,
        "response_modes": 30,
        "applies_to": 24,
        "content": 76,
        "operator_instruction": 58,
        "template_hint": 58,
        "source": 36,
        "source_type": 18,
        "source_ref": 52,
        "source_quote": 70,
        "source_date": 16,
        "notes": 44,
        "entity_id": 28,
        "canonical_value": 34,
        "aliases": 64,
        "key": 30,
        "value": 50,
        "description": 76,
    }
    for ws in wb.worksheets:
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        for cell in ws[1]:
            cell.fill = fill
            cell.font = font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        headers = sheet_headers(ws, [])
        for idx, header in enumerate(headers, start=1):
            ws.column_dimensions[get_column_letter(idx)].width = default_widths.get(header, 24)


def rows_as_dicts(ws: Any, required_headers: list[str]) -> tuple[list[str], list[dict[str, Any]]]:
    headers = sheet_headers(ws, required_headers)
    rows = []
    for row_idx in range(2, ws.max_row + 1):
        row = {h: "" for h in headers}
        empty = True
        for col_idx, header in enumerate(headers, start=1):
            value = ws.cell(row=row_idx, column=col_idx).value
            row[header] = "" if value is None else str(value)
            if row[header]:
                empty = False
        if not empty:
            row["_row_idx"] = row_idx
            rows.append(row)
    return headers, rows


def write_row(ws: Any, headers: list[str], row_idx: int, data: dict[str, Any]) -> None:
    for col_idx, header in enumerate(headers, start=1):
        ws.cell(row=row_idx, column=col_idx).value = data.get(header, "")


def append_row(ws: Any, headers: list[str], data: dict[str, Any]) -> None:
    row_idx = ws.max_row + 1
    write_row(ws, headers, row_idx, data)


def ask_multiline_text() -> str:
    print("Paste text. Finish with a single line: END")
    lines = []
    while True:
        line = input()
        if line.strip() == "END":
            break
        lines.append(line)
    return clean_text("\n".join(lines))


def call_ollama(prompt: str) -> dict[str, Any] | None:
    url = os.getenv("OLLAMA_URL", "http://localhost:11434/api/generate")
    model = os.getenv("OLLAMA_MODEL", "gpt-oss:120b-cloud")
    payload = json.dumps(
        {"model": model, "prompt": prompt, "stream": False, "options": {"temperature": 0.1}},
        ensure_ascii=False,
    ).encode("utf-8")
    req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=180) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except Exception as exc:
        print(f"ERROR: LLM request failed: {exc}")
        return None

    raw = str(data.get("response", "")).strip()
    if not raw:
        return None

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass

    match = re.search(r"\{.*\}", raw, flags=re.S)
    if not match:
        return None
    try:
        return json.loads(match.group(0))
    except json.JSONDecodeError:
        return None


def llm_extract(text: str, source_ref: str, source_type: str) -> dict[str, Any] | None:
    prompt = f"""
Ты помогаешь заполнять Knowledge Base для e-mail обработки входящих сообщений info@mooon.by.

Задача: превратить источник в структурированные KB-записи.

Главное:
- извлекай только стабильные знания;
- не придумывай политики, которых нет в источнике;
- не классифицируй по одному слову;
- не сохраняй live-данные как факты Knowledge Base;
- расписание, текущие сеансы, цены, наличие мест, live-репертуар не являются KB;
- если источник содержит динамические данные, не переноси их в content;
- если из источника нет стабильного знания, верни пустые items и dictionary;
- source_quote должен быть точной короткой цитатой из источника, если цитата реально есть;
- если точной цитаты нет, source_quote оставь пустым;
- каждая отдельная мысль/правило/инструкция — отдельный item;
- верни только JSON без markdown.

category = о чём знание, lower_snake_case. Примеры:
refunds, tickets, payment, schedule, loyalty, certificates, partners, legal, documents, marketing, technical, system_noise, general.

type = как использовать знание:
rule, faq, operator_instruction, response_template, regulation, historical_case.

response_modes разрешены только:
answer; ask_clarifying_question; handoff_to_operator; no_reply/ignore.

applies_to примеры:
customer; partner; government; counterparty; system; spam; operator.

Формат JSON:
{{
  "items": [
    {{
      "id": "english_stable_id",
      "category": "lower_snake_case_category",
      "type": "rule|faq|operator_instruction|response_template|regulation|historical_case",
      "title": "русский заголовок",
      "priority": "1-10",
      "topic_keywords": "ключи через ;",
      "entity_keywords": "ключи через ;",
      "customer_need_keywords": "ключи через ;",
      "response_modes": "answer; ask_clarifying_question; handoff_to_operator; no_reply/ignore",
      "applies_to": "customer|partner|government|counterparty|system|spam|operator",
      "content": "стабильное знание на русском",
      "operator_instruction": "что должен делать оператор, если применимо",
      "template_hint": "подсказка для будущего ответа, если применимо",
      "source_quote": "точная короткая цитата из источника, если есть",
      "source_date": "дата источника, если явно есть",
      "notes": "важные ограничения"
    }}
  ],
  "dictionary": [
    {{
      "entity_id": "english_id",
      "type": "brand|venue|service|intent|term|document",
      "canonical_value": "нормализованное значение",
      "aliases": "синонимы через ;",
      "notes": "заметки"
    }}
  ],
  "warnings": ["если часть источника была динамической/неподходящей — кратко напиши здесь"]
}}

source_ref: {source_ref}
source_type: {source_type}

Текст источника:
{text[:18000]}
""".strip()
    return call_ollama(prompt)


def validate_response_modes(value: str) -> str:
    modes = split_keywords(value)
    valid = [mode for mode in modes if mode in ALLOWED_RESPONSE_MODES]
    return "; ".join(valid) if valid else "handoff_to_operator"


def validate_item(raw: dict[str, Any], source_ref: str, source_type: str) -> dict[str, str]:
    title = str(raw.get("title") or "Знание из добавленного источника").strip()
    category = str(raw.get("category") or "general").strip().lower()
    category = re.sub(r"[^a-z0-9_\-]+", "_", category).strip("_") or "general"

    item_type = str(raw.get("type") or "rule").strip()
    if item_type not in ALLOWED_ITEM_TYPES:
        item_type = "rule"

    item_id = str(raw.get("id") or slugify(f"{category}_{title}")).strip().lower()
    item_id = re.sub(r"[^a-zA-Z0-9_\-]+", "_", item_id).strip("_").lower() or slugify(title)

    priority = str(raw.get("priority") or "5").strip()
    if not re.fullmatch(r"10|[1-9]", priority):
        priority = "5"

    source_quote = clean_text(str(raw.get("source_quote") or ""))
    if len(source_quote) > 700:
        source_quote = source_quote[:700].rstrip() + "..."

    return {
        "id": item_id,
        "enabled": "yes",
        "category": category,
        "type": item_type,
        "title": title,
        "priority": priority,
        "topic_keywords": str(raw.get("topic_keywords") or "").strip(),
        "entity_keywords": str(raw.get("entity_keywords") or "").strip(),
        "customer_need_keywords": str(raw.get("customer_need_keywords") or "").strip(),
        "response_modes": validate_response_modes(str(raw.get("response_modes") or "handoff_to_operator")),
        "applies_to": str(raw.get("applies_to") or "customer").strip(),
        "content": clean_text(str(raw.get("content") or "")),
        "operator_instruction": clean_text(str(raw.get("operator_instruction") or "")),
        "template_hint": clean_text(str(raw.get("template_hint") or "")),
        "source": source_ref,
        "source_type": str(raw.get("source_type") or source_type).strip(),
        "source_ref": str(raw.get("source_ref") or source_ref).strip(),
        "source_quote": source_quote,
        "source_date": str(raw.get("source_date") or "").strip(),
        "notes": clean_text(str(raw.get("notes") or "")),
    }


def validate_dictionary(raw: dict[str, Any]) -> dict[str, str]:
    canonical = str(raw.get("canonical_value") or "").strip()
    entity_id = str(raw.get("entity_id") or slugify(canonical or "entity", "entity")).strip().lower()
    entity_id = re.sub(r"[^a-zA-Z0-9_\-]+", "_", entity_id).strip("_") or "entity"
    return {
        "entity_id": entity_id,
        "type": str(raw.get("type") or "term").strip(),
        "canonical_value": canonical,
        "aliases": str(raw.get("aliases") or "").strip(),
        "notes": clean_text(str(raw.get("notes") or "")),
    }


def item_match_text(item: dict[str, Any]) -> str:
    return " ".join([
        str(item.get("category", "")),
        str(item.get("title", "")),
        str(item.get("topic_keywords", "")),
        str(item.get("entity_keywords", "")),
        str(item.get("customer_need_keywords", "")),
        str(item.get("content", "")),
    ])


def find_similar_item(existing: list[dict[str, Any]], item: dict[str, str]) -> dict[str, Any] | None:
    for row in existing:
        if str(row.get("id", "")).strip().lower() == item["id"].lower():
            return row

    best = None
    best_score = 0.0
    candidate_text = item_match_text(item)
    for row in existing:
        score = similarity(candidate_text, item_match_text(row))
        if score > best_score:
            best = row
            best_score = score
    if best and best_score >= 0.42:
        return best
    return None


def merge_item(old: dict[str, Any], new: dict[str, str], headers: list[str]) -> dict[str, str]:
    merged = {h: str(old.get(h, "")) for h in headers}

    for field in ["category", "type", "title", "priority", "source_type", "source_date"]:
        if not merged.get(field) and new.get(field):
            merged[field] = new[field]

    for field in ["topic_keywords", "entity_keywords", "customer_need_keywords", "response_modes", "applies_to", "source", "source_ref"]:
        merged[field] = join_unique(merged.get(field, ""), new.get(field, ""))

    for field in ["content", "operator_instruction", "template_hint", "source_quote"]:
        new_value = clean_text(new.get(field, ""))
        old_value = clean_text(merged.get(field, ""))
        if new_value and new_value not in old_value:
            merged[field] = clean_text((old_value + "\n\n" + new_value).strip())

    note = f"Updated {datetime.now().strftime('%Y-%m-%d %H:%M')} from {new.get('source_ref', '')}"
    merged["notes"] = clean_text((merged.get("notes", "") + "\n" + new.get("notes", "") + "\n" + note).strip())
    if not merged.get("enabled"):
        merged["enabled"] = "yes"
    return merged


def upsert_dictionary(ws: Any, entries: list[dict[str, str]]) -> tuple[int, int]:
    headers, rows = rows_as_dicts(ws, KB_DICTIONARY_HEADERS)
    added = 0
    updated = 0

    for entry in entries:
        if not entry.get("canonical_value"):
            continue
        found = None
        for row in rows:
            if str(row.get("entity_id", "")).lower() == entry["entity_id"].lower() or str(row.get("canonical_value", "")).lower() == entry["canonical_value"].lower():
                found = row
                break
        if found:
            row_idx = int(found["_row_idx"])
            merged = {h: str(found.get(h, "")) for h in headers}
            merged["aliases"] = join_unique(merged.get("aliases", ""), entry.get("aliases", ""))
            if entry.get("notes") and entry["notes"] not in merged.get("notes", ""):
                merged["notes"] = clean_text((merged.get("notes", "") + "\n" + entry["notes"]).strip())
            write_row(ws, headers, row_idx, merged)
            updated += 1
        else:
            append_row(ws, headers, entry)
            rows.append({**entry, "_row_idx": ws.max_row})
            added += 1
    return added, updated


def print_item_preview(item: dict[str, str]) -> None:
    print(f"  id={item.get('id')}")
    print(f"  category={item.get('category')} type={item.get('type')} priority={item.get('priority')}")
    print(f"  title={item.get('title')}")
    content = clean_text(item.get("content", ""))
    if content:
        print(f"  content={content[:240]}{'...' if len(content) > 240 else ''}")
    quote = clean_text(item.get("source_quote", ""))
    if quote:
        print(f"  quote={quote[:240]}{'...' if len(quote) > 240 else ''}")
    print(f"  source_ref={item.get('source_ref')}")


def apply_to_workbook(data: dict[str, Any], source_ref: str, source_type: str) -> None:
    wb = ensure_workbook()
    ws_items = wb["KB_ITEMS"]
    ws_dict = wb["KB_DICTIONARY"]

    raw_items = data.get("items") or []
    raw_dict = data.get("dictionary") or []
    warnings = data.get("warnings") or []

    if not isinstance(raw_items, list) or not isinstance(raw_dict, list):
        print("ERROR: invalid LLM JSON structure")
        return

    items = [validate_item(x, source_ref, source_type) for x in raw_items if isinstance(x, dict)]
    dict_entries = [validate_dictionary(x) for x in raw_dict if isinstance(x, dict)]

    if warnings:
        print("\nWarnings:")
        for warning in warnings:
            print(f"- {warning}")

    if not items and not dict_entries:
        print("No stable KB entries extracted. Workbook was not changed.")
        return

    item_headers, existing = rows_as_dicts(ws_items, KB_ITEMS_HEADERS)
    print("\nPlanned changes:")
    planned: list[tuple[str, dict[str, str], dict[str, Any] | None]] = []
    for item in items:
        similar = find_similar_item(existing, item)
        if similar:
            print(f"\n- update KB_ITEMS: {similar.get('id')} <- {item.get('title')}")
            print_item_preview(item)
            planned.append(("update", item, similar))
        else:
            print(f"\n- add KB_ITEMS: {item.get('id')} — {item.get('title')}")
            print_item_preview(item)
            planned.append(("add", item, None))

    if dict_entries:
        print(f"\n- upsert KB_DICTIONARY: {len(dict_entries)} item(s)")

    answer = input("\nApply changes? y/n: ").strip().lower()
    if answer != "y":
        print("Cancelled")
        return

    if KB_PATH.exists():
        shutil.copy2(KB_PATH, BACKUP_PATH)
        print(f"backup={BACKUP_PATH}")

    added_items = 0
    updated_items = 0
    for action, item, similar in planned:
        if action == "update" and similar:
            row_idx = int(similar["_row_idx"])
            merged = merge_item(similar, item, item_headers)
            write_row(ws_items, item_headers, row_idx, merged)
            updated_items += 1
        else:
            append_row(ws_items, item_headers, item)
            added_items += 1

    added_dict, updated_dict = upsert_dictionary(ws_dict, dict_entries)
    style_workbook(wb)
    wb.save(KB_PATH)
    print(f"saved={KB_PATH}")
    print(f"items_added={added_items} items_updated={updated_items} dict_added={added_dict} dict_updated={updated_dict}")


def disable_kb_item() -> None:
    wb = ensure_workbook()
    ws_items = wb["KB_ITEMS"]
    headers, rows = rows_as_dicts(ws_items, KB_ITEMS_HEADERS)
    item_id = input("KB item id to disable: ").strip()
    if not item_id:
        print("ERROR: empty id")
        return

    found = None
    for row in rows:
        if str(row.get("id", "")).strip().lower() == item_id.lower():
            found = row
            break

    if not found:
        print(f"ERROR: item not found: {item_id}")
        return

    reason = input("Reason/note: ").strip()
    print(f"\nPlanned changes:\n- disable KB_ITEMS: {found.get('id')} — {found.get('title')}")
    answer = input("\nApply changes? y/n: ").strip().lower()
    if answer != "y":
        print("Cancelled")
        return

    if KB_PATH.exists():
        shutil.copy2(KB_PATH, BACKUP_PATH)
        print(f"backup={BACKUP_PATH}")

    row_idx = int(found["_row_idx"])
    updated = {h: str(found.get(h, "")) for h in headers}
    updated["enabled"] = "no"
    note = f"Disabled {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    if reason:
        note += f": {reason}"
    updated["notes"] = clean_text((updated.get("notes", "") + "\n" + note).strip())
    write_row(ws_items, headers, row_idx, updated)
    style_workbook(wb)
    wb.save(KB_PATH)
    print(f"saved={KB_PATH}")
    print("items_disabled=1")


def choose_source() -> tuple[str, str, str] | None:
    print("\nKB Update")
    print("1. Paste text")
    print("2. Read file")
    print("3. Read URL")
    print("4. Disable KB item")
    print("q. Exit")
    choice = input("Choice: ").strip().lower()

    if choice == "1":
        text = ask_multiline_text()
        return text, "manual_text", "manual_text"
    if choice == "2":
        path = Path(input("File path: ").strip().strip('"')).expanduser()
        text = read_file(path)
        source_ref = str(path)
        return text, source_ref, infer_source_type(source_ref)
    if choice == "3":
        url = input("URL: ").strip()
        text = read_url(url)
        return text, url, "url"
    if choice == "4":
        disable_kb_item()
        return None
    if choice == "q":
        raise SystemExit(0)
    raise RuntimeError("Unknown choice")


def main() -> None:
    print(f"project_root={PROJECT_ROOT}")
    print(f"kb={KB_PATH}")

    chosen = choose_source()
    if chosen is None:
        return

    text, source_ref, source_type = chosen
    if not text.strip():
        print("ERROR: empty input")
        return

    print("extracting=llm")
    data = llm_extract(text, source_ref, source_type)
    if data is None:
        print("ERROR: LLM did not return valid JSON. Workbook was not changed.")
        return

    apply_to_workbook(data, source_ref, source_type)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\nCancelled")
    except Exception as exc:
        print(f"ERROR: {exc}")
        raise
